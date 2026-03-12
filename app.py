"""
Marketing ROI & Footfall Analytics System — Flask Application
"""

import calendar
import os
import io
import re
from datetime import date, datetime
from functools import wraps

import mysql.connector
from mysql.connector import pooling
import json
import numpy as np
import pandas as pd
from flask import (Flask, render_template, request, redirect,
                   url_for, session, flash)
from werkzeug.security import check_password_hash
from werkzeug.utils import secure_filename

from config import Config

# ── App setup ────────────────────────────────────────────────
app = Flask(__name__)
print("Hello World!")
app.config.from_object(Config)

MONTH_ORDER = [
    'January', 'February', 'March', 'April', 'May', 'June',
    'July', 'August', 'September', 'October', 'November', 'December'
]


# ── Custom Jinja2 filter for number formatting ───────────────
@app.template_filter('numfmt')
def number_format(value, decimals=0):
    try:
        return '{:,.{d}f}'.format(float(value), d=int(decimals))
    except (ValueError, TypeError):
        return value


# ── Database helpers ─────────────────────────────────────────
db_pool = None


def init_pool():
    global db_pool
    db_pool = pooling.MySQLConnectionPool(
        pool_name="app_pool",
        pool_size=5,
        host=Config.MYSQL_HOST,
        user=Config.MYSQL_USER,
        password=Config.MYSQL_PASSWORD,
        database=Config.MYSQL_DATABASE,
        port=Config.MYSQL_PORT,
    )


def get_db():
    global db_pool
    if db_pool is None:
        init_pool()
    return db_pool.get_connection()


# ── Auth decorator ───────────────────────────────────────────
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated


# ── Helper: latest month in monthly_stock_summary ────────────
def _latest_summary_month(cursor):
    cursor.execute("""
        SELECT month_name, year FROM monthly_stock_summary
        ORDER BY year DESC,
               FIELD(month_name,'January','February','March','April',
                     'May','June','July','August','September',
                     'October','November','December') DESC
        LIMIT 1
    """)
    row = cursor.fetchone()
    return row if row else None


# ══════════════════════════════════════════════════════════════
#  ROUTES
# ══════════════════════════════════════════════════════════════

# ── Login / Logout ───────────────────────────────────────────
@app.route('/', methods=['GET'])
def index():
    return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')

        if not username or not password:
            flash('Please enter both username and password.', 'error')
            return render_template('login.html')

        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        try:
            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            user = cursor.fetchone()
        finally:
            cursor.close()
            conn.close()

        if user and check_password_hash(user['password_hash'], password):
            session['user_id'] = user['id']
            session['username'] = user['username']
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password.', 'error')

    return render_template('login.html')


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


# ── Admin Dashboard (Products & Overview) ────────────────────
@app.route('/dashboard')
@login_required
def dashboard():
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # All products
        cursor.execute("SELECT * FROM products ORDER BY name")
        products = cursor.fetchall()

        # Latest month
        latest = _latest_summary_month(cursor)
        latest_month = latest['month_name'] if latest else None
        latest_year = latest['year'] if latest else None

        # Stock overview based on latest monthly closing
        stock_overview = []
        if latest_month:
            cursor.execute("""
                SELECT p.id, p.name, p.unit,
                    COALESCE(m.opening_qty, 0) + COALESCE(m.received_qty, 0) AS total_in,
                    COALESCE(m.sold_qty, 0)   AS total_sold,
                    COALESCE(m.closing_qty, 0) AS in_stock,
                    COALESCE(m.closing_value, 0) AS stock_value,
                    CASE WHEN COALESCE(m.closing_qty, 0) > 0
                         THEN ROUND(m.closing_value / m.closing_qty, 2)
                         ELSE 0 END AS price_per_unit
                FROM products p
                LEFT JOIN monthly_stock_summary m
                    ON p.id = m.product_id
                   AND m.month_name = %s AND m.year = %s
                ORDER BY p.name
            """, (latest_month, latest_year))
            stock_overview = cursor.fetchall()

            # Adjust for entries added AFTER the latest summary month
            month_num = MONTH_ORDER.index(latest_month) + 1
            last_day = calendar.monthrange(latest_year, month_num)[1]
            cutoff = date(latest_year, month_num, last_day)

            cursor.execute("""
                SELECT product_id,
                    COALESCE(SUM(CASE WHEN entry_type IN ('OPENING','RECEIVED')
                                      THEN qty ELSE 0 END), 0) AS new_in,
                    COALESCE(SUM(CASE WHEN entry_type = 'SOLD'
                                      THEN qty ELSE 0 END), 0) AS new_sold,
                    COALESCE(SUM(CASE WHEN entry_type IN ('OPENING','RECEIVED')
                                      THEN total_value ELSE 0 END), 0) AS new_in_val,
                    COALESCE(SUM(CASE WHEN entry_type = 'SOLD'
                                      THEN total_value ELSE 0 END), 0) AS new_sold_val
                FROM stock_entries
                WHERE entry_date > %s
                GROUP BY product_id
            """, (cutoff,))
            adjustments = {r['product_id']: r for r in cursor.fetchall()}

            for item in stock_overview:
                adj = adjustments.get(item['id'], {})
                item['total_in'] = float(item['total_in'] or 0) + float(adj.get('new_in', 0))
                item['total_sold'] = float(item['total_sold'] or 0) + float(adj.get('new_sold', 0))
                item['in_stock'] = float(item['in_stock'] or 0) + float(adj.get('new_in', 0)) - float(adj.get('new_sold', 0))
                item['stock_value'] = float(item['stock_value'] or 0) + float(adj.get('new_in_val', 0)) - float(adj.get('new_sold_val', 0))
                item['price_per_unit'] = round(item['stock_value'] / item['in_stock'], 2) if item['in_stock'] > 0 else 0
        else:
            # No monthly data yet — build overview purely from entries
            cursor.execute("""
                SELECT p.id, p.name, p.unit,
                    COALESCE(SUM(CASE WHEN se.entry_type IN ('OPENING','RECEIVED')
                                      THEN se.qty ELSE 0 END), 0) AS total_in,
                    COALESCE(SUM(CASE WHEN se.entry_type = 'SOLD'
                                      THEN se.qty ELSE 0 END), 0) AS total_sold,
                    COALESCE(SUM(CASE WHEN se.entry_type IN ('OPENING','RECEIVED')
                                      THEN se.qty ELSE 0 END), 0) -
                    COALESCE(SUM(CASE WHEN se.entry_type = 'SOLD'
                                      THEN se.qty ELSE 0 END), 0) AS in_stock,
                    COALESCE(SUM(CASE WHEN se.entry_type IN ('OPENING','RECEIVED')
                                      THEN se.total_value ELSE 0 END), 0) -
                    COALESCE(SUM(CASE WHEN se.entry_type = 'SOLD'
                                      THEN se.total_value ELSE 0 END), 0) AS stock_value,
                    0 AS price_per_unit
                FROM products p
                LEFT JOIN stock_entries se ON p.id = se.product_id
                GROUP BY p.id, p.name, p.unit
                ORDER BY p.name
            """)
            stock_overview = cursor.fetchall()
            for item in stock_overview:
                item['price_per_unit'] = round(float(item['stock_value']) / float(item['in_stock']), 2) if float(item['in_stock']) > 0 else 0

        # Available months for history
        cursor.execute("""
            SELECT DISTINCT month_name, year
            FROM monthly_stock_summary
            ORDER BY year,
                     FIELD(month_name,'January','February','March','April',
                           'May','June','July','August','September',
                           'October','November','December')
        """)
        available_months = cursor.fetchall()

    finally:
        cursor.close()
        conn.close()

    return render_template('admin.html',
                           products=products,
                           stock_overview=stock_overview,
                           available_months=available_months,
                           latest_month=latest_month,
                           latest_year=latest_year)


# ── Add product ──────────────────────────────────────────────
@app.route('/add-product', methods=['POST'])
@login_required
def add_product():
    name = request.form.get('product_name', '').strip()
    unit = request.form.get('unit', '').strip().lower()
    custom_unit = request.form.get('custom_unit', '').strip().lower()
    price = request.form.get('price_per_unit', '0').strip()

    # Use custom unit when "+ New Unit" was selected
    if unit == '__new__':
        unit = custom_unit

    if not name or not unit:
        flash('Product name and unit are required.', 'error')
        return redirect(url_for('dashboard'))

    try:
        price = round(float(price), 2) if price else 0.0
    except ValueError:
        price = 0.0

    conn = get_db()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT IGNORE INTO products (name, unit, price_per_unit) VALUES (%s, %s, %s)",
            (name, unit, price)
        )
        conn.commit()
        if cursor.rowcount:
            flash(f'Product "{name}" added successfully!', 'success')
        else:
            flash(f'Product "{name}" already exists.', 'error')
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for('dashboard'))


# ── Delete product ───────────────────────────────────────────
@app.route('/delete-product/<int:pid>', methods=['POST'])
@login_required
def delete_product(pid):
    conn = get_db()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM products WHERE id = %s", (pid,))
        conn.commit()
        flash('Product deleted.', 'success')
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for('dashboard'))


# ── Clear ALL Database Data (DANGER) ─────────────────────────
@app.route('/clear-database', methods=['POST'])
@login_required
def clear_database():
    """Clear all data from the database (except users). USE WITH EXTREME CAUTION!"""
    conn = get_db()
    cursor = conn.cursor()
    try:
        # Delete all data from tables (in reverse order of dependencies)
        cursor.execute("SET FOREIGN_KEY_CHECKS = 0")  # Temporarily disable FK checks
        cursor.execute("TRUNCATE TABLE stock_entries")
        cursor.execute("TRUNCATE TABLE monthly_stock_summary")
        cursor.execute("TRUNCATE TABLE marketing_spend")
        cursor.execute("TRUNCATE TABLE products")
        cursor.execute("SET FOREIGN_KEY_CHECKS = 1")  # Re-enable FK checks
        conn.commit()
        flash('✅ Database cleared successfully! All data has been deleted (users preserved).', 'success')
    except Exception as e:
        conn.rollback()
        flash(f'❌ Error clearing database: {str(e)}', 'error')
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for('dashboard'))


# ── Stock Movement Log ───────────────────────────────────────
@app.route('/stock-log')
@login_required
def stock_log():
    product_filter = request.args.get('product', 'all')
    month_filter = request.args.get('month', 'all')
    year_filter = request.args.get('year', 'all')

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("SELECT id, name FROM products ORDER BY name")
        products = cursor.fetchall()
        
        # Get distinct months and years
        cursor.execute("""
            SELECT DISTINCT 
                DATE_FORMAT(entry_date, '%M') AS month_name,
                MONTH(entry_date) AS month_num,
                YEAR(entry_date) AS year
            FROM stock_entries
            ORDER BY year DESC, month_num DESC
        """)
        date_data = cursor.fetchall()
        months = sorted(set(d['month_name'] for d in date_data), key=lambda x: datetime.strptime(x, '%B').month)
        years = sorted(set(d['year'] for d in date_data), reverse=True)

        # Build query with filters
        query = """
            SELECT se.*, p.name AS product_name, p.unit
            FROM stock_entries se
            JOIN products p ON se.product_id = p.id
            WHERE 1=1
        """
        params = []
        
        if product_filter != 'all' and product_filter.isdigit():
            query += " AND se.product_id = %s"
            params.append(int(product_filter))
            
        if month_filter != 'all':
            query += " AND DATE_FORMAT(se.entry_date, '%M') = %s"
            params.append(month_filter)
            
        if year_filter != 'all':
            query += " AND YEAR(se.entry_date) = %s"
            params.append(int(year_filter))
        
        query += " ORDER BY se.entry_date DESC, se.id DESC"
        
        cursor.execute(query, params)
        entries = cursor.fetchall()
    finally:
        cursor.close()
        conn.close()

    return render_template('stock_log.html',
                           entries=entries,
                           products=products,
                           product_filter=product_filter,
                           month_filter=month_filter,
                           year_filter=year_filter,
                           months=months,
                           years=years)


# ── Add stock entry ──────────────────────────────────────────
@app.route('/add-stock-entry', methods=['POST'])
@login_required
def add_stock_entry():
    product_id = request.form.get('product_id', '')
    entry_type = request.form.get('entry_type', '')
    qty = request.form.get('qty', '')
    price_per_unit = request.form.get('price_per_unit', '')
    note = request.form.get('note', '').strip()
    entry_date = request.form.get('entry_date', '')

    # Validate
    if not all([product_id, entry_type, qty, price_per_unit, entry_date]):
        flash('All fields are required.', 'error')
        return redirect(url_for('stock_log'))

    if entry_type not in ('OPENING', 'RECEIVED', 'SOLD'):
        flash('Invalid entry type.', 'error')
        return redirect(url_for('stock_log'))

    try:
        qty_f = float(qty)
        ppu_f = float(price_per_unit)
        if qty_f <= 0 or ppu_f < 0:
            raise ValueError
    except ValueError:
        flash('Quantity must be positive and price non-negative.', 'error')
        return redirect(url_for('stock_log'))

    total_value = round(qty_f * ppu_f, 2)

    conn = get_db()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO stock_entries
            (entry_date, product_id, entry_type, qty, price_per_unit, total_value, note)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (entry_date, int(product_id), entry_type, qty_f, ppu_f, total_value, note or None))
        conn.commit()
        flash('Stock entry added!', 'success')
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for('stock_log'))


# ── Delete stock entry ───────────────────────────────────────
@app.route('/delete-entry/<int:eid>', methods=['POST'])
@login_required
def delete_entry(eid):
    conn = get_db()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM stock_entries WHERE id = %s", (eid,))
        conn.commit()
        flash('Entry deleted.', 'success')
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for('stock_log'))


@app.route('/delete-entries-bulk', methods=['POST'])
@login_required
def delete_entries_bulk():
    """Delete multiple stock entries at once."""
    entry_ids = request.form.getlist('entry_ids[]')
    
    if not entry_ids:
        flash('No entries selected.', 'error')
        return redirect(url_for('stock_log'))
    
    conn = get_db()
    cursor = conn.cursor()
    try:
        placeholders = ','.join(['%s'] * len(entry_ids))
        cursor.execute(f"DELETE FROM stock_entries WHERE id IN ({placeholders})", entry_ids)
        conn.commit()
        flash(f'{len(entry_ids)} entries deleted successfully.', 'success')
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for('stock_log'))


# ── Monthly History ──────────────────────────────────────────
@app.route('/monthly-history')
@login_required
def monthly_history():
    selected_month = request.args.get('month', '')
    selected_year = request.args.get('year', '')

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # Available months
        cursor.execute("""
            SELECT DISTINCT month_name, year
            FROM monthly_stock_summary
            ORDER BY year,
                     FIELD(month_name,'January','February','March','April',
                           'May','June','July','August','September',
                           'October','November','December')
        """)
        available_months = cursor.fetchall()

        stock_data = []
        marketing_summary = {}
        total_sales = 0
        chart_data = {}
        all_months_sales = []
        insights = {}

        # Fetch sales totals for ALL months (for comparison)
        cursor.execute("""
            SELECT m.month_name, m.year,
                   COALESCE(SUM(m.sold_value), 0) AS total_sales
            FROM monthly_stock_summary m
            GROUP BY m.month_name, m.year
            ORDER BY m.year,
                     FIELD(m.month_name,'January','February','March','April',
                           'May','June','July','August','September',
                           'October','November','December')
        """)
        all_months_sales = cursor.fetchall()

        if selected_month and selected_year:
            # Stock data for selected month
            cursor.execute("""
                SELECT p.name, p.unit,
                       m.opening_qty, m.opening_value,
                       m.received_qty, m.received_value,
                       m.sold_qty, m.sold_value,
                       m.closing_qty, m.closing_value,
                       m.remarks
                FROM monthly_stock_summary m
                JOIN products p ON m.product_id = p.id
                WHERE m.month_name = %s AND m.year = %s
                ORDER BY p.name
            """, (selected_month, int(selected_year)))
            stock_data = cursor.fetchall()

            total_sales = sum(float(r['sold_value'] or 0) for r in stock_data)

            # Build insights: best/worst product, top product by qty
            if stock_data:
                best_prod = max(stock_data, key=lambda r: float(r['sold_value'] or 0))
                worst_prod = min(stock_data, key=lambda r: float(r['sold_value'] or 0))
                top_qty = max(stock_data, key=lambda r: float(r['sold_qty'] or 0))
                insights = {
                    'best_product': best_prod['name'],
                    'best_value': float(best_prod['sold_value'] or 0),
                    'worst_product': worst_prod['name'],
                    'worst_value': float(worst_prod['sold_value'] or 0),
                    'top_qty_product': top_qty['name'],
                    'top_qty': float(top_qty['sold_qty'] or 0),
                }

            # Build chart data for this month's products
            chart_data = {
                'product_names': [r['name'] for r in stock_data],
                'sold_qty':      [float(r['sold_qty'] or 0) for r in stock_data],
                'sold_value':    [float(r['sold_value'] or 0) for r in stock_data],
                'opening_qty':   [float(r['opening_qty'] or 0) for r in stock_data],
                'received_qty':  [float(r['received_qty'] or 0) for r in stock_data],
                'closing_qty':   [float(r['closing_qty'] or 0) for r in stock_data],
                'all_months':    [r['month_name'] + ' ' + str(r['year']) for r in all_months_sales],
                'all_months_sales': [float(r['total_sales'] or 0) for r in all_months_sales],
            }

            # Marketing summary for the month
            cursor.execute("""
                SELECT
                    COUNT(*)                     AS days,
                    COALESCE(SUM(digital_spend), 0)   AS total_digital,
                    COALESCE(SUM(print_spend), 0)     AS total_print,
                    COALESCE(SUM(outdoor_spend), 0)   AS total_outdoor,
                    COALESCE(SUM(total_spend), 0)     AS total_spend,
                    COALESCE(SUM(customer_walk_ins), 0) AS total_walkins,
                    COALESCE(SUM(sales_amount), 0)    AS total_mkt_sales
                FROM marketing_spend
                WHERE month_name = %s
            """, (selected_month,))
            marketing_summary = cursor.fetchone() or {}

            # Marketing channel data for pie chart
            if marketing_summary:
                chart_data['mkt_channels'] = ['Digital', 'Print', 'Outdoor']
                chart_data['mkt_values'] = [
                    float(marketing_summary.get('total_digital', 0) or 0),
                    float(marketing_summary.get('total_print', 0) or 0),
                    float(marketing_summary.get('total_outdoor', 0) or 0),
                ]

    finally:
        cursor.close()
        conn.close()

    return render_template('monthly_history.html',
                           available_months=available_months,
                           stock_data=stock_data,
                           marketing_summary=marketing_summary,
                           total_sales=total_sales,
                           selected_month=selected_month,
                           selected_year=selected_year,
                           insights=insights,
                           all_months_sales=all_months_sales,
                           chart_data_json=json.dumps(chart_data))


# ══════════════════════════════════════════════════════════════
#  UPLOAD DATA
# ══════════════════════════════════════════════════════════════

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls', 'pdf', 'docx'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def _allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def _parse_file_to_df(filepath):
    """Read uploaded file into a pandas DataFrame."""
    ext = filepath.rsplit('.', 1)[1].lower()
    if ext == 'csv':
        return pd.read_csv(filepath)
    elif ext in ('xlsx', 'xls'):
        return pd.read_excel(filepath, engine='openpyxl')
    elif ext == 'pdf':
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        text = '\n'.join(page.extract_text() or '' for page in reader.pages)
        return _text_to_df(text)
    elif ext == 'docx':
        from docx import Document
        doc = Document(filepath)
        # Try to find a table first
        if doc.tables:
            table = doc.tables[0]
            data = [[c.text.strip() for c in row.cells] for row in table.rows]
            if len(data) > 1:
                return pd.DataFrame(data[1:], columns=data[0])
        # Fallback: parse text
        text = '\n'.join(p.text for p in doc.paragraphs)
        return _text_to_df(text)
    return pd.DataFrame()


def _text_to_df(text):
    """Best-effort: parse tabular text (comma/tab separated) into a DataFrame."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return pd.DataFrame()
    # Detect separator
    sep = ',' if ',' in lines[0] else '\t' if '\t' in lines[0] else ','
    return pd.read_csv(io.StringIO('\n'.join(lines)), sep=sep)


def _safe_float(val):
    try:
        return float(str(val).strip().replace(',', '')) if pd.notna(val) else 0.0
    except (ValueError, TypeError):
        return 0.0


def _safe_int(val):
    try:
        return int(float(str(val).strip().replace(',', ''))) if pd.notna(val) else 0
    except (ValueError, TypeError):
        return 0


def _detect_data_type(cols_lower):
    """Auto-detect whether the uploaded data is stock or marketing based on columns."""
    # Marketing indicators
    marketing_keywords = ['digital', 'print_spend', 'outdoor', 'walk_in', 'customer_walk',
                          'digital_spend', 'outdoor_spend', 'sales_amount']
    # Stock indicators
    stock_keywords = ['item', 'product', 'opening', 'received', 'closing', 'sold_qty',
                      'opening_qty', 'closing_qty', 'received_qty']

    marketing_score = sum(1 for kw in marketing_keywords if any(kw in c for c in cols_lower))
    stock_score = sum(1 for kw in stock_keywords if any(kw in c for c in cols_lower))

    if marketing_score >= 2:
        return 'marketing'
    if stock_score >= 2:
        return 'stock'
    return None


def _import_stock_df(df, year, cursor):
    """Import stock DataFrame into monthly_stock_summary + products + stock_entries."""
    # Normalise column names
    col_map = {}
    for c in df.columns:
        cl = c.strip().lower().replace(' ', '_')
        if 'month' in cl:
            col_map[c] = 'Month'
        elif 'item' in cl or 'product' in cl:
            col_map[c] = 'Item'
        elif 'unit' in cl:
            col_map[c] = 'Unit'
        elif 'opening' in cl and 'qty' in cl:
            col_map[c] = 'Opening_Qty'
        elif 'opening' in cl and 'val' in cl:
            col_map[c] = 'Opening_Value'
        elif 'received' in cl and 'qty' in cl:
            col_map[c] = 'Received_Qty'
        elif 'received' in cl and 'val' in cl:
            col_map[c] = 'Received_Value'
        elif 'sold' in cl and 'qty' in cl:
            col_map[c] = 'Sold_Qty'
        elif 'sold' in cl and 'val' in cl:
            col_map[c] = 'Sold_Value'
        elif 'closing' in cl and 'qty' in cl:
            col_map[c] = 'Closing_Qty'
        elif 'closing' in cl and 'val' in cl:
            col_map[c] = 'Closing_Value'
        elif 'remark' in cl:
            col_map[c] = 'Remarks'
    df = df.rename(columns=col_map)

    required = ['Month', 'Item']
    if not all(c in df.columns for c in required):
        return 0, 'Missing required columns: Month, Item'

    count = 0
    for _, row in df.iterrows():
        month = str(row.get('Month', '')).strip()
        item = str(row.get('Item', '')).strip()
        unit = str(row.get('Unit', 'nos')).strip().lower()
        if not month or not item:
            continue

        # Ensure product exists
        cursor.execute("INSERT IGNORE INTO products (name, unit) VALUES (%s, %s)", (item, unit))
        cursor.execute("SELECT id FROM products WHERE name = %s", (item,))
        pid = cursor.fetchone()['id']

        cursor.execute("""
            INSERT INTO monthly_stock_summary
            (month_name, year, product_id, opening_qty, opening_value,
             received_qty, received_value, sold_qty, sold_value,
             closing_qty, closing_value, remarks)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                opening_qty=VALUES(opening_qty), opening_value=VALUES(opening_value),
                received_qty=VALUES(received_qty), received_value=VALUES(received_value),
                sold_qty=VALUES(sold_qty), sold_value=VALUES(sold_value),
                closing_qty=VALUES(closing_qty), closing_value=VALUES(closing_value),
                remarks=VALUES(remarks)
        """, (
            month, int(year), pid,
            _safe_float(row.get('Opening_Qty')), _safe_float(row.get('Opening_Value')),
            _safe_float(row.get('Received_Qty')), _safe_float(row.get('Received_Value')),
            _safe_float(row.get('Sold_Qty')), _safe_float(row.get('Sold_Value')),
            _safe_float(row.get('Closing_Qty')), _safe_float(row.get('Closing_Value')),
            str(row.get('Remarks', '')).strip() if pd.notna(row.get('Remarks')) else None,
        ))

        # Also create stock entries
        month_map = {'January':1,'February':2,'March':3,'April':4,'May':5,'June':6,
                     'July':7,'August':8,'September':9,'October':10,'November':11,'December':12}
        mnum = month_map.get(month)
        if mnum:
            yr = int(year)
            last_day = calendar.monthrange(yr, mnum)[1]
            for etype, qty_col, val_col, day, note in [
                ('OPENING', 'Opening_Qty', 'Opening_Value', 1, f'{month} opening stock'),
                ('RECEIVED', 'Received_Qty', 'Received_Value', 15, f'{month} received stock'),
                ('SOLD', 'Sold_Qty', 'Sold_Value', last_day, f'{month} sales'),
            ]:
                qty = _safe_float(row.get(qty_col))
                val = _safe_float(row.get(val_col))
                if qty > 0:
                    ppu = round(val / qty, 2) if qty else 0
                    cursor.execute("""
                        INSERT INTO stock_entries
                        (entry_date, product_id, entry_type, qty, price_per_unit, total_value, note)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                    """, (date(yr, mnum, day), pid, etype, qty, ppu, val, note))
        count += 1
    return count, None


def _import_marketing_df(df, cursor):
    """Import marketing DataFrame into marketing_spend."""
    col_map = {}
    for c in df.columns:
        cl = c.strip().lower().replace(' ', '_')
        if cl == 'date':
            col_map[c] = 'Date'
        elif 'month' in cl:
            col_map[c] = 'Month'
        elif 'day' in cl and 'type' in cl:
            col_map[c] = 'Day_Type'
        elif 'digital' in cl:
            col_map[c] = 'Digital_Spend'
        elif 'print' in cl:
            col_map[c] = 'Print_Spend'
        elif 'outdoor' in cl:
            col_map[c] = 'Outdoor_Spend'
        elif 'total' in cl and 'spend' in cl:
            col_map[c] = 'Total_Spend'
        elif 'walk' in cl or 'customer' in cl:
            col_map[c] = 'Customer_Walk_Ins'
        elif 'sales' in cl or 'amount' in cl:
            col_map[c] = 'Sales_Amount'
    df = df.rename(columns=col_map)

    if 'Date' not in df.columns or 'Month' not in df.columns:
        return 0, 'Missing required columns: Date, Month'

    count = 0
    for _, row in df.iterrows():
        date_str = str(row.get('Date', '')).strip()
        month = str(row.get('Month', '')).strip()
        if not date_str or not month:
            continue
        try:
            # Try multiple date formats
            spend_date = None
            for fmt in ('%d-%m-%Y', '%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y'):
                try:
                    spend_date = datetime.strptime(date_str, fmt).date()
                    break
                except ValueError:
                    continue
            if not spend_date:
                continue
        except Exception:
            continue

        digital = _safe_float(row.get('Digital_Spend'))
        print_s = _safe_float(row.get('Print_Spend'))
        outdoor = _safe_float(row.get('Outdoor_Spend'))
        total = _safe_float(row.get('Total_Spend')) or (digital + print_s + outdoor)

        cursor.execute("""
            INSERT INTO marketing_spend
            (spend_date, month_name, day_type, digital_spend, print_spend,
             outdoor_spend, total_spend, customer_walk_ins, sales_amount)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            spend_date, month,
            str(row.get('Day_Type', '')).strip(),
            digital, print_s, outdoor, total,
            _safe_int(row.get('Customer_Walk_Ins')),
            _safe_float(row.get('Sales_Amount')),
        ))
        count += 1
    return count, None


@app.route('/upload-data', methods=['GET', 'POST'])
@login_required
def upload_data():
    upload_result = None

    if request.method == 'POST':
        file = request.files.get('file')
        year = request.form.get('year', str(datetime.now().year))

        if not file or file.filename == '':
            flash('Please select a file.', 'error')
            return redirect(url_for('upload_data'))

        if not _allowed_file(file.filename):
            flash('Unsupported file format. Use CSV, Excel, PDF or Word.', 'error')
            return redirect(url_for('upload_data'))

        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        try:
            df = _parse_file_to_df(filepath)
            # Drop completely empty rows
            df = df.dropna(how='all')

            if df.empty:
                upload_result = {'success': False, 'message': 'Could not extract any data from the file.', 'details': []}
            else:
                # Auto-detect data type from column names
                cols_lower = [c.strip().lower().replace(' ', '_') for c in df.columns]
                data_type = _detect_data_type(cols_lower)

                conn = get_db()
                cursor = conn.cursor(dictionary=True)
                try:
                    if data_type == 'stock':
                        count, err = _import_stock_df(df, year, cursor)
                        if err:
                            upload_result = {'success': False, 'message': err, 'details': [f'Columns found: {list(df.columns)}']}
                        else:
                            conn.commit()
                            upload_result = {
                                'success': True,
                                'message': f'Successfully imported {count} stock records.',
                                'details': [
                                    f'File: {filename}',
                                    f'Data type detected: Stock & Sales',
                                    f'Rows processed: {count}',
                                    f'Year: {year}',
                                    'Products, monthly summaries, and stock entries updated.'
                                ]
                            }
                    elif data_type == 'marketing':
                        count, err = _import_marketing_df(df, cursor)
                        if err:
                            upload_result = {'success': False, 'message': err, 'details': [f'Columns found: {list(df.columns)}']}
                        else:
                            conn.commit()
                            upload_result = {
                                'success': True,
                                'message': f'Successfully imported {count} marketing records.',
                                'details': [
                                    f'File: {filename}',
                                    f'Data type detected: Marketing Spend',
                                    f'Rows processed: {count}'
                                ]
                            }
                    else:
                        upload_result = {
                            'success': False,
                            'message': 'Could not identify the data type automatically.',
                            'details': [
                                f'Columns found: {list(df.columns)}',
                                f'Rows: {len(df)}',
                                'Tip: For stock data, include columns like Month, Item, Opening_Qty, Sold_Qty, etc.',
                                'Tip: For marketing data, include columns like Date, Digital_Spend, Sales_Amount, etc.',
                                'Make sure the file is clean with no null values or errors.'
                            ]
                        }
                finally:
                    cursor.close()
                    conn.close()
        except Exception as e:
            upload_result = {'success': False, 'message': f'Error processing file: {str(e)}', 'details': []}
        finally:
            # Clean up uploaded file
            if os.path.exists(filepath):
                os.remove(filepath)

    return render_template('upload.html', upload_result=upload_result)


# ══════════════════════════════════════════════════════════════
#  PREDICTIONS
# ══════════════════════════════════════════════════════════════

def _build_prediction_model():
    """Train a linear regression model from individual daily marketing records.

    Training on daily rows (50+ data points) instead of monthly aggregates
    gives the model much more variation and prevents wild extrapolation
    when the user enters budgets outside the historical monthly range.
    """
    from sklearn.linear_model import LinearRegression

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # --- Individual daily records for model training ---
        cursor.execute("""
            SELECT digital_spend, print_spend, outdoor_spend,
                   customer_walk_ins, sales_amount
            FROM marketing_spend
            WHERE sales_amount > 0
        """)
        daily_rows = cursor.fetchall()

        # --- Monthly aggregates for chart display & recommendations ---
        cursor.execute("""
            SELECT month_name,
                   SUM(digital_spend) AS digital, SUM(print_spend) AS print_s,
                   SUM(outdoor_spend) AS outdoor, SUM(total_spend) AS total,
                   SUM(customer_walk_ins) AS walkins, SUM(sales_amount) AS sales,
                   COUNT(*) AS days_count
            FROM marketing_spend
            GROUP BY month_name
            HAVING COUNT(*) >= 2
            ORDER BY FIELD(month_name,'January','February','March','April',
                           'May','June','July','August','September',
                           'October','November','December')
        """)
        monthly_rows = cursor.fetchall()
    finally:
        cursor.close()
        conn.close()

    if len(daily_rows) < 3:
        return None, None, monthly_rows, None

    X = np.array([[float(r['digital_spend']), float(r['print_spend']),
                   float(r['outdoor_spend'])] for r in daily_rows])
    y_sales = np.array([float(r['sales_amount']) for r in daily_rows])
    y_walkins = np.array([float(r['customer_walk_ins']) for r in daily_rows])

    model_sales = LinearRegression().fit(X, y_sales)
    model_walkins = LinearRegression().fit(X, y_walkins)

    # Pre-compute fallback ratios (sales per Rs of total spend)
    # Used when model extrapolation goes negative for very small budgets
    total_spends = X.sum(axis=1)
    fallback = {
        'sales_ratio': float(y_sales.sum() / total_spends.sum()),
        'walkins_ratio': float(y_walkins.sum() / total_spends.sum()),
        'n_daily': len(daily_rows),
        'avg_days_per_month': np.mean([float(r['days_count']) for r in monthly_rows]) if monthly_rows else 10,
    }

    return model_sales, model_walkins, monthly_rows, fallback


@app.route('/predictions', methods=['GET', 'POST'])
@login_required
def predictions():
    prediction = None
    prediction_json = '{}'
    
    # Get current month and next month based on latest data
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # Get latest month from monthly_stock_summary
        cursor.execute("""
            SELECT month_name, year 
            FROM monthly_stock_summary 
            ORDER BY year DESC, 
                FIELD(month_name,'January','February','March','April','May','June',
                      'July','August','September','October','November','December') DESC
            LIMIT 1
        """)
        latest_month_data = cursor.fetchone()
    finally:
        cursor.close()
        conn.close()
    
    current_month = latest_month_data['month_name'] if latest_month_data else 'Unknown'
    current_year = latest_month_data['year'] if latest_month_data else datetime.now().year
    
    # Calculate next month
    if current_month != 'Unknown':
        month_names = ['January', 'February', 'March', 'April', 'May', 'June',
                       'July', 'August', 'September', 'October', 'November', 'December']
        current_idx = month_names.index(current_month)
        next_idx = (current_idx + 1) % 12
        next_month = month_names[next_idx]
        next_year = current_year + 1 if next_idx == 0 else current_year
    else:
        next_month = 'Unknown'
        next_year = current_year

    if request.method == 'POST':
        digital = float(request.form.get('digital_spend', 0))
        print_s = float(request.form.get('print_spend', 0))
        outdoor = float(request.form.get('outdoor_spend', 0))
        total_budget = digital + print_s + outdoor

        model_sales, model_walkins, hist_rows, fallback = _build_prediction_model()

        if model_sales is None:
            flash('Not enough data to make predictions. Upload more marketing data.', 'error')
            return render_template('predictions.html', prediction=None, prediction_json='{}',
                                   current_month=current_month, current_year=current_year,
                                   next_month=next_month, next_year=next_year)

        # Convert monthly budget to daily average for the model
        # (model is trained on individual daily records)
        avg_days = fallback['avg_days_per_month']
        daily_input = np.array([[digital / avg_days, print_s / avg_days, outdoor / avg_days]])

        # Predict daily sales/walk-ins, then scale to full month (30 days)
        days_in_month = 30
        pred_sales = float(model_sales.predict(daily_input)[0]) * days_in_month
        pred_walkins = float(model_walkins.predict(daily_input)[0]) * days_in_month

        # Fallback: if model extrapolates to negative (budget far outside training range),
        # use historical average ratio instead
        if pred_sales <= 0 and total_budget > 0:
            pred_sales = total_budget * fallback['sales_ratio']
        if pred_walkins <= 0 and total_budget > 0:
            pred_walkins = total_budget * fallback['walkins_ratio']

        pred_sales = max(0, round(pred_sales))
        pred_walkins = max(0, int(round(pred_walkins)))
        roi = round(pred_sales / total_budget, 1) if total_budget > 0 else 0

        # Historical data for charts
        hist_months = [r['month_name'] for r in hist_rows]
        hist_sales = [float(r['sales']) for r in hist_rows]

        # Confidence note based on daily data points
        n = fallback['n_daily']
        if n >= 30:
            confidence_note = f'Good amount of historical data ({n} daily records) — predictions should be fairly reliable.'
        elif n >= 10:
            confidence_note = f'Moderate data available ({n} daily records). More months will improve accuracy.'
        else:
            confidence_note = f'Limited data ({n} records). Upload more months for better predictions.'

        # Stock recommendations
        stock_recs = _get_stock_recommendations()

        # Smart recommendations
        recommendations = []
        avg_digital = np.mean([float(r['digital']) for r in hist_rows])
        avg_print = np.mean([float(r['print_s']) for r in hist_rows])
        avg_outdoor = np.mean([float(r['outdoor']) for r in hist_rows])

        if digital > avg_digital * 1.3:
            recommendations.append({
                'icon': '📱', 'color': '#4facfe',
                'title': 'High Digital Spend',
                'detail': f'Your digital budget (Rs {digital:,.0f}) is significantly above your average (Rs {avg_digital:,.0f}). '
                          'Digital marketing often gives the best ROI for hardware stores through Google Maps and local search ads.'
            })
        elif digital < avg_digital * 0.7:
            recommendations.append({
                'icon': '📱', 'color': '#f5576c',
                'title': 'Consider Increasing Digital Spend',
                'detail': f'Digital spend (Rs {digital:,.0f}) is below average (Rs {avg_digital:,.0f}). '
                          'More digital presence can significantly boost walk-ins.'
            })

        if print_s > avg_print * 1.3:
            recommendations.append({
                'icon': '📰', 'color': '#11998e',
                'title': 'High Print Spend',
                'detail': f'Print budget (Rs {print_s:,.0f}) is above average (Rs {avg_print:,.0f}). '
                          'Local newspaper ads work well for reaching hardware buyers in your area.'
            })

        if outdoor > avg_outdoor * 1.3:
            recommendations.append({
                'icon': '🪧', 'color': '#f7971e',
                'title': 'Good Outdoor Visibility',
                'detail': f'Outdoor spend (Rs {outdoor:,.0f}) is above average. Banners and signage help attract passing customers.'
            })

        if total_budget > 0 and roi > 5:
            recommendations.append({
                'icon': '🏆', 'color': '#11998e',
                'title': 'Excellent Expected ROI',
                'detail': f'Expected ROI of {roi}x means every Rs 1 spent could return Rs {roi} in sales. This is a strong budget allocation.'
            })
        elif total_budget > 0 and roi < 3:
            recommendations.append({
                'icon': '⚠️', 'color': '#f5576c',
                'title': 'Consider Budget Reallocation',
                'detail': f'Expected ROI of {roi}x is moderate. Consider shifting more budget to the channel that historically performs best.'
            })

        # Weekend tip
        recommendations.append({
            'icon': '📅', 'color': '#764ba2',
            'title': 'Focus on Weekends',
            'detail': 'Historical data shows weekends have significantly higher walk-ins and sales. '
                      'Increase marketing intensity on Fridays and Saturdays for better results.'
        })

        prediction = {
            'total_budget': total_budget,
            'predicted_sales': round(pred_sales),
            'predicted_walkins': pred_walkins,
            'roi': roi,
            'data_points': n,
            'confidence_note': confidence_note,
            'stock_recommendations': stock_recs,
            'recommendations': recommendations,
        }

        prediction_json = json.dumps({
            'digital': digital, 'print_val': print_s, 'outdoor': outdoor,
            'predicted_sales': round(pred_sales),
            'hist_months': hist_months,
            'hist_sales': hist_sales,
        })

    return render_template('predictions.html',
                           prediction=prediction,
                           prediction_json=prediction_json,
                           current_month=current_month,
                           current_year=current_year,
                           next_month=next_month,
                           next_year=next_year)


def _get_stock_recommendations():
    """Get stock purchase recommendations based on recent sales trends."""
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # Average sold qty per product across all months
        cursor.execute("""
            SELECT p.name, p.unit,
                   AVG(m.sold_qty) AS avg_sold,
                   (SELECT m2.closing_qty FROM monthly_stock_summary m2
                    WHERE m2.product_id = p.id
                    ORDER BY m2.year DESC,
                             FIELD(m2.month_name,'January','February','March','April',
                                   'May','June','July','August','September',
                                   'October','November','December') DESC
                    LIMIT 1) AS current_closing
            FROM products p
            JOIN monthly_stock_summary m ON m.product_id = p.id
            GROUP BY p.id, p.name, p.unit
            ORDER BY p.name
        """)
        rows = cursor.fetchall()
    finally:
        cursor.close()
        conn.close()

    recs = []
    for r in rows:
        avg_sold = float(r['avg_sold'] or 0)
        closing = float(r['current_closing'] or 0)
        # Recommend purchasing 1.3x average sold minus current closing
        needed = round(avg_sold * 1.3 - closing)
        if needed < 0:
            needed = 0
        remark = ''
        if closing < avg_sold * 0.5:
            remark = 'Stock critically low — order urgently!'
        elif closing < avg_sold:
            remark = 'Below average demand — restock soon'
        elif needed == 0:
            remark = 'Stock sufficient for next month'
        else:
            remark = f'Need ~{needed} more to cover 130% of avg demand'

        recs.append({
            'name': r['name'],
            'unit': r['unit'],
            'avg_sold': avg_sold,
            'current_closing': closing,
            'recommended_purchase': max(0, needed),
            'remark': remark,
        })
    return recs


# ══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    app.run(debug=True, port=5000)
