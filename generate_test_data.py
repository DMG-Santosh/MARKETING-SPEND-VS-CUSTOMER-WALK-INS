"""
Generate synthetic test data for upload feature testing
"""
import pandas as pd
from docx import Document

# Create Excel file for June stock data
stock_june = {
    'Month': ['June']*9,
    'Item': ['Glass', 'Plywood', 'Hardwares', 'Beeding', 'Laminates', 'Doors', 'Cement', 'Steel Rods', 'Paints'],
    'Unit': ['SFT', 'SFT', 'NOS', 'RFT', 'NOS', 'NOS', 'BAG', 'TON', 'LTR'],
    'Opening_Qty': [335, 315, 2067, 2250, 475, 250, 530, 9, 180],
    'Opening_Value': [518580, 264600, 216750, 361600, 283900, 300000, 159000, 270000, 216000],
    'Received_Qty': [300, 270, 1800, 2200, 550, 240, 700, 18, 250],
    'Received_Value': [464400, 226800, 189000, 352000, 341000, 288000, 210000, 540000, 300000],
    'Sold_Qty': [210, 205, 1450, 2700, 430, 190, 600, 20, 200],
    'Sold_Value': [324996, 172200, 152250, 432000, 266900, 228000, 180000, 600000, 240000],
    'Closing_Qty': [425, 380, 2417, 1750, 595, 300, 630, 7, 230],
    'Closing_Value': [658284, 319200, 253500, 281600, 358000, 360000, 189000, 210000, 276000],
    'Remarks': ['Monsoon prep', 'Waterproof', 'Rainy essentials', 'Lower demand', 'Moisture resistant', 'Teak wood', 'Monsoon stock', 'Construction slow', 'Waterproof range']
}
df_stock = pd.DataFrame(stock_june)
df_stock.to_excel('test_data/stock_june_2026.xlsx', index=False, engine='openpyxl')
print('✓ Created stock_june_2026.xlsx')

# Create Word document for marketing data
doc = Document()
doc.add_heading('Marketing Spend Data - June 2026', 0)
table = doc.add_table(rows=1, cols=9)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
headers = ['Date', 'Month', 'Day_Type', 'Digital_Spend', 'Print_Spend', 'Outdoor_Spend', 'Total_Spend', 'Customer_Walk_Ins', 'Sales_Amount']
for i, h in enumerate(headers):
    hdr_cells[i].text = h

marketing_data = [
    ['01-06-2026', 'June', 'Weekday', '1400', '580', '380', '2360', '44', '52800'],
    ['02-06-2026', 'June', 'Weekday', '1350', '560', '360', '2270', '42', '50400'],
    ['03-06-2026', 'June', 'Weekday', '1450', '600', '400', '2450', '47', '56400'],
    ['04-06-2026', 'June', 'Weekday', '1500', '620', '420', '2540', '49', '58800'],
    ['05-06-2026', 'June', 'Weekday', '1550', '650', '450', '2650', '52', '62400'],
    ['06-06-2026', 'June', 'Weekend', '2600', '1050', '850', '4500', '96', '115200'],
    ['07-06-2026', 'June', 'Weekend', '2700', '1100', '900', '4700', '101', '121200'],
    ['08-06-2026', 'June', 'Weekday', '1450', '600', '400', '2450', '48', '57600'],
]
for data_row in marketing_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(data_row):
        row_cells[i].text = val
doc.save('test_data/marketing_june_2026.docx')
print('✓ Created marketing_june_2026.docx')

print('\n✅ All test files created successfully!')
print('\nTest files created in test_data folder:')
print('  - stock_may_2026.csv (CSV format)')
print('  - marketing_may_2026.csv (CSV format)')
print('  - stock_june_2026.xlsx (Excel format)')
print('  - marketing_june_2026.docx (Word format)')
